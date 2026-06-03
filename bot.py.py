import os
import random
from docx import Document

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    ApplicationBuilder,
    CommandHandler,
    CallbackQueryHandler,
    MessageHandler,
    ContextTypes,
    filters,
)

TOKEN = "8795576766:AAH0wi-nRxHbzMuGJoYCLn_JLSjjNxaACpE"

user_data = {}


# ================= RASM MAP =================
def extract_images(doc, img_dir="images"):
    os.makedirs(img_dir, exist_ok=True)
    img_map = {}

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            path = os.path.join(img_dir, f"{rel.rId}.png")
            with open(path, "wb") as f:
                f.write(rel.target_part.blob)
            img_map[rel.rId] = path

    return img_map


# ================= CELL PARSE =================
def parse_cell(cell, img_map):
    parts = []

    for p in cell.paragraphs:
        for run in p.runs:
            if run.text:
                parts.append({"type": "text", "value": run.text})

            blips = run.element.xpath(".//a:blip")
            for blip in blips:
                rId = blip.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                )
                if rId in img_map:
                    parts.append({"type": "image", "value": img_map[rId]})

    return parts


# ================= WORD O‘QISH =================
def load_test(file_path):
    doc = Document(file_path)
    img_map = extract_images(doc)

    questions = []

    for table in doc.tables:
        cells = []
        for row in table.rows:
            for cell in row.cells:
                cells.append(cell)

        if len(cells) < 3:
            continue

        question_parts = parse_cell(cells[0], img_map)

        options = []
        for cell in cells[1:]:
            parts = parse_cell(cell, img_map)
            options.append(parts)

        correct = options[0]
        random.shuffle(options)
        correct_index = options.index(correct)

        questions.append({
            "question": question_parts,
            "options": options,
            "correct": correct_index
        })

    return questions


# ================= START =================
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "🤖 made by saidmaruf\n\n📄 Word fayl yuboring (.docx)"
    )


# ================= FILE =================
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    file = update.message.document

    if not file.file_name.endswith(".docx"):
        await update.message.reply_text("❗ Faqat .docx yuboring")
        return

    file_path = f"{file.file_id}.docx"

    new_file = await context.bot.get_file(file.file_id)
    await new_file.download_to_drive(file_path)

    questions = load_test(file_path)

    if not questions:
        await update.message.reply_text("❌ Test topilmadi")
        return

    total = len(questions)
    part = total // 4

    sections = [
        questions[0:part],
        questions[part:part*2],
        questions[part*2:part*3],
        questions[part*3:total]
    ]

    user_data[update.message.from_user.id] = {
        "sections": sections
    }

    keyboard = [
        [InlineKeyboardButton(f"{i+1}-bo‘lim", callback_data=f"section_{i}")]
        for i in range(4)
    ]

    await update.message.reply_text(
        "Bo‘limni tanlang:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


# ================= BO‘LIM =================
async def choose_section(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user = user_data[query.from_user.id]
    idx = int(query.data.split("_")[1])

    user["questions"] = user["sections"][idx]
    user["index"] = 0
    user["wrong"] = []

    await send_question(query, context)


# ================= SAVOL =================
async def send_question(query, context):
    user = user_data[query.from_user.id]
    q = user["questions"][user["index"]]

    text = ""
    image_path = None

    for part in q["question"]:
        if part["type"] == "text":
            text += part["value"]
        elif part["type"] == "image":
            image_path = part["value"]

    if image_path:
        await query.message.reply_photo(photo=open(image_path, "rb"), caption=text)
    else:
        await query.message.reply_text(text)

    # ===== VARIANT =====
    for i, opt in enumerate(q["options"]):
        opt_text = ""
        opt_image = None

        for part in opt:
            if part["type"] == "text":
                opt_text += part["value"]
            elif part["type"] == "image":
                opt_image = part["value"]

        msg = f"{i+1}. {opt_text}"

        if opt_image:
            await query.message.reply_photo(photo=open(opt_image, "rb"), caption=msg)
        else:
            await query.message.reply_text(msg)

    keyboard = [
        [InlineKeyboardButton(f"{i+1}", callback_data=f"ans_{i}")]
        for i in range(len(q["options"]))
    ]

    await query.message.reply_text(
        "Javobni tanlang:",
        reply_markup=InlineKeyboardMarkup(keyboard)
    )


# ================= JAVOB =================
async def answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()

    user = user_data[query.from_user.id]
    q = user["questions"][user["index"]]

    selected = int(query.data.split("_")[1])
    correct = q["correct"]

    correct_text = ""
    for part in q["options"][correct]:
        if part["type"] == "text":
            correct_text += part["value"]

    if selected == correct:
        await query.message.reply_text(
            f"✅ To‘g‘ri!\n\n✔ To‘g‘ri javob:\n{correct+1}. {correct_text}"
        )
    else:
        await query.message.reply_text(
            f"❌ Xato!\n\n✔ To‘g‘ri javob:\n{correct+1}. {correct_text}"
        )

        if q not in user["wrong"]:
            user["wrong"].append(q)

    user["index"] += 1

    if user["index"] >= len(user["questions"]):
        if user["wrong"]:
            user["questions"] = user["wrong"].copy()
            user["wrong"].clear()
            user["index"] = 0
            random.shuffle(user["questions"])

            await query.message.reply_text("🔁 Xato savollar qayta beriladi")
        else:
            await query.message.reply_text("🎉 Bo‘lim tugadi!")
            return

    await send_question(query, context)


# ================= MAIN =================
def main():
    print("🤖 Bot ishga tushdi... (made by saidmaruf)")

    app = ApplicationBuilder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(CallbackQueryHandler(choose_section, pattern="section_"))
    app.add_handler(CallbackQueryHandler(answer, pattern="ans_"))

    app.run_polling()


if __name__ == "__main__":
    main()